#!/usr/bin/env python3
"""Render cl-data.yaml + cl-format.yaml → cover-letter.docx

Deterministic cover letter document generator. Reads structured YAML data
and a config file, produces a consistently formatted Word document that
visually matches the CV.
"""

import yaml
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers (shared logic with render_cv — kept self-contained for simplicity)
# ---------------------------------------------------------------------------

def load_yaml(path):
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def hex_to_rgb(hex_color):
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_spacing(paragraph, before=0, after=0, line=None):
    pPr = paragraph._p.get_or_add_pPr()
    # Remove existing spacing element to avoid duplicates
    for existing in pPr.findall(qn("w:spacing")):
        pPr.remove(existing)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))
    if line is not None:
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)


def _add_bottom_border(paragraph, color="000000", size="6"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), size)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _add_hyperlink(paragraph, url, text, color_hex, font=None, size=None, bold=False):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if font:
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), font)
        rFonts.set(qn("w:hAnsi"), font)
        rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(int(size * 2)))
        rPr.append(szCs)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color_hex.lstrip("#"))
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run_el.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run_el.append(t)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def _run(paragraph, text, *, font, size, bold=False, italic=False, color=None):
    r = paragraph.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return r


# ---------------------------------------------------------------------------
# Main renderer
# ---------------------------------------------------------------------------

def render_cl(data, config, output_path):
    """Build a cover letter .docx from structured *data* using *config*."""
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    pg = config["page"]
    if pg["size"] == "A4":
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    else:
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.94)
    m = pg["margins"]
    section.top_margin = Cm(m["top"])
    section.bottom_margin = Cm(m["bottom"])
    section.left_margin = Cm(m["left"])
    section.right_margin = Cm(m["right"])

    typo = config["typography"]
    heading_font = typo.get("heading_font", typo.get("font", "Calibri"))
    body_font = typo.get("body_font", typo.get("font", "Calibri"))
    hdr_cfg = config["header"]
    hdr = data["header"]
    primary = hex_to_rgb(hdr_cfg.get("name_color", config["colors"]["primary"]))
    link_hex = config["colors"]["link"]
    line_sp = int(typo["line_spacing"] * 240)
    hdr_align = (WD_ALIGN_PARAGRAPH.LEFT if hdr_cfg.get("alignment", "center") == "left"
                 else WD_ALIGN_PARAGRAPH.CENTER)

    # ── Header (matches CV) ──
    # Name
    np = doc.add_paragraph()
    np.alignment = hdr_align
    _set_spacing(np, after=0)
    name_text = (hdr["name"].upper() if hdr_cfg["name_case"] == "upper"
                 else hdr["name"].title() if hdr_cfg["name_case"] == "title"
                 else hdr["name"])
    _run(np, name_text, font=heading_font, size=hdr_cfg["name_font_size"],
         bold=hdr_cfg.get("name_bold", True), color=primary)

    # Title
    if hdr_cfg.get("show_title", True) and hdr.get("title"):
        tp = doc.add_paragraph()
        tp.alignment = hdr_align
        _set_spacing(tp, after=0)
        _run(tp, hdr["title"], font=body_font,
             size=hdr_cfg.get("title_font_size", typo["body_size"]))

    # Contact
    contact = []
    for key in ("phone", "email"):
        if hdr.get(key):
            contact.append(hdr[key])
    if contact:
        cp = doc.add_paragraph()
        cp.alignment = hdr_align
        _set_spacing(cp, after=0)
        _run(cp, " | ".join(contact), font=body_font,
             size=hdr_cfg.get("contact_font_size", typo["body_size"]))

    # Links
    has_links = hdr.get("linkedin") or hdr.get("github")
    if has_links:
        lp = doc.add_paragraph()
        lp.alignment = hdr_align
        _set_spacing(lp, after=0)
        links_size = hdr_cfg.get("links_font_size", typo["body_size"])
        if hdr.get("linkedin"):
            display = hdr["linkedin"].replace("https://", "").replace("http://", "").rstrip("/")
            _add_hyperlink(lp, hdr["linkedin"], display, link_hex, font=body_font, size=links_size)
        if hdr.get("linkedin") and hdr.get("github"):
            _run(lp, " | ", font=body_font, size=links_size)
        if hdr.get("github"):
            display = hdr["github"].replace("https://", "").replace("http://", "").rstrip("/")
            _add_hyperlink(lp, hdr["github"], display, link_hex, font=body_font, size=links_size)

    # Relocation
    if hdr_cfg.get("show_relocation", True) and hdr.get("relocation"):
        rp = doc.add_paragraph()
        rp.alignment = hdr_align
        _set_spacing(rp, after=0)
        _run(rp, hdr["relocation"], font=body_font,
             size=hdr_cfg.get("relocation_font_size", typo["body_size"]),
             bold=hdr_cfg.get("relocation_bold", True))

    # Horizontal rule under header
    hr = doc.add_paragraph()
    _set_spacing(hr, before=60, after=120)
    _add_bottom_border(hr)

    # ── Date ──
    date_cfg = config.get("date", {})
    dp = doc.add_paragraph()
    align = date_cfg.get("alignment", "left")
    dp.alignment = (WD_ALIGN_PARAGRAPH.LEFT if align == "left"
                    else WD_ALIGN_PARAGRAPH.RIGHT)
    _set_spacing(dp, after=0)
    _run(dp, data["date"], font=body_font, size=typo["body_size"])

    # Empty paragraph for visible blank line between date and salutation
    blank = doc.add_paragraph()
    _set_spacing(blank, after=0, line=int(typo["line_spacing"] * 240))
    _run(blank, "", font=body_font, size=typo["body_size"])

    # ── Salutation ──
    sp = doc.add_paragraph()
    _set_spacing(sp, after=120)
    _run(sp, data["salutation"], font=body_font, size=typo["body_size"])

    # ── Body paragraphs ──
    body_cfg = config.get("body", {})
    para_after = body_cfg.get("paragraph_spacing_after", 120)
    for para_text in data["paragraphs"]:
        pp = doc.add_paragraph()
        _set_spacing(pp, after=para_after, line=line_sp)
        _run(pp, para_text.strip(), font=body_font, size=typo["body_size"])

    # ── Sign-off ──
    so = data["sign_off"]
    so_cfg = config.get("sign_off", {})

    # Closing line
    clp = doc.add_paragraph()
    _set_spacing(clp, before=120, after=240)
    _run(clp, so.get("closing", so_cfg.get("closing_text", "Kind regards,")),
         font=body_font, size=typo["body_size"])

    # Name
    snp = doc.add_paragraph()
    _set_spacing(snp, after=0)
    _run(snp, so["name"], font=body_font, size=typo["body_size"],
         bold=so_cfg.get("name_bold", True),
         color=hex_to_rgb(so_cfg.get("name_color", config["colors"]["primary"])))

    # Title, email, phone
    for field in ("title", "email", "phone"):
        if so.get(field):
            fp = doc.add_paragraph()
            _set_spacing(fp, after=0)
            _run(fp, so[field], font=body_font, size=typo["body_size"])

    doc.save(str(output_path))
    return str(output_path)
