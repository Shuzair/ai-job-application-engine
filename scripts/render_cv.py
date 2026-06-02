#!/usr/bin/env python3
"""Render cv-data.yaml + cv-format.yaml → cv.docx

Deterministic CV document generator. Reads structured YAML data and a
config file, produces a consistently formatted Word document.
"""

import yaml
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def load_yaml(path):
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def hex_to_rgb(hex_color):
    """'#2E74B5' → RGBColor."""
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_spacing(paragraph, before=0, after=0, line=None):
    """Set paragraph spacing in twips."""
    pPr = paragraph._p.get_or_add_pPr()
    # Remove existing spacing element to avoid duplicates
    for existing in pPr.findall(qn("w:spacing")):
        pPr.remove(existing)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))
    if line is not None:
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)


def _add_bottom_border(paragraph, color="000000", size="6"):
    """Add a horizontal rule (bottom border) to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), size)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _add_hyperlink(paragraph, url, text, color_hex, font=None, size=None, bold=False):
    """Insert a clickable hyperlink run into *paragraph*."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if font:
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), font)
        rFonts.set(qn("w:hAnsi"), font)
        rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(int(size * 2)))
        rPr.append(szCs)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color_hex.lstrip("#"))
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run_el.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run_el.append(t)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def _run(paragraph, text, *, font, size, bold=False, italic=False, color=None):
    """Add a formatted run and return it."""
    r = paragraph.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def _add_inline_text(paragraph, text, *, font, size, color=None):
    """Add text to paragraph, parsing **...** as bold runs."""
    import re
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            _run(paragraph, part[2:-2], font=font, size=size, bold=True, color=color)
        elif part:
            _run(paragraph, part, font=font, size=size, color=color)


def _remove_table_borders(table):
    """Remove all borders from a table, making it invisible."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


# ---------------------------------------------------------------------------
# Section renderers
# ---------------------------------------------------------------------------

def _heading(doc, title, cfg, key=None):
    sec = cfg["sections"]
    if key:
        title = sec.get("headings", {}).get(key, title)
    typo = cfg["typography"]
    heading_font = typo.get("heading_font", typo.get("font", "Calibri"))
    text = title.upper() if sec["heading_case"] == "upper" else (
        title.title() if sec["heading_case"] == "title" else title
    )
    p = doc.add_paragraph()
    _set_spacing(p, before=200, after=40)
    primary = hex_to_rgb(cfg["colors"]["primary"])
    _run(p, text, font=heading_font, size=sec["heading_size"],
         bold=sec.get("heading_bold", True), color=primary)
    if sec.get("heading_underline", True):
        _add_bottom_border(p)


def _render_summary(doc, data, cfg, _w):
    if not data.get("summary"):
        return
    _heading(doc, "Professional Summary", cfg, key="summary")
    typo = cfg["typography"]
    body_font = typo.get("body_font", typo.get("font", "Calibri"))
    p = doc.add_paragraph()
    _set_spacing(p, after=60, line=int(typo["line_spacing"] * 240))
    _run(p, data["summary"].strip(), font=body_font, size=typo["body_size"])


def _render_skills(doc, data, cfg, _w):
    if not data.get("skills"):
        return
    _heading(doc, "Skills", cfg, key="skills")
    typo = cfg["typography"]
    body_font = typo.get("body_font", typo.get("font", "Calibri"))
    sep = cfg.get("skills", {}).get("separator", " · ")
    for grp in data["skills"]:
        p = doc.add_paragraph()
        _set_spacing(p, after=20, line=int(typo["line_spacing"] * 240))
        _run(p, f"{grp['category']}:", font=body_font,
             size=typo["body_size"], bold=True)
        _run(p, f" {sep.join(grp['items'])}", font=body_font,
             size=typo["body_size"])


def _render_experience(doc, data, cfg, usable_w):
    if not data.get("experience"):
        return
    _heading(doc, "Professional Experience", cfg, key="experience")
    typo = cfg["typography"]
    body_font = typo.get("body_font", typo.get("font", "Calibri"))
    exp_cfg = cfg.get("experience", {})
    line_sp = int(typo["line_spacing"] * 240)

    sep = cfg.get("experience", {}).get("separator", " | ")

    for i, role in enumerate(data["experience"]):
        # Single line: Title | Company | Location  →  date_start – date_end
        tp = doc.add_paragraph()
        _set_spacing(tp, before=120 if i else 0, after=0)
        tp.paragraph_format.tab_stops.add_tab_stop(usable_w, WD_TAB_ALIGNMENT.RIGHT)
        _run(tp, role["title"], font=body_font, size=typo["body_size"], bold=True)

        # Inline company and location
        company = role["company"]
        _run(tp, sep, font=body_font, size=typo["body_size"])
        _run(tp, company, font=body_font, size=typo["body_size"])

        hq = role.get("hq") or ""
        loc = role.get("location") or ""
        wm = role.get("work_model") or ""
        if hq or loc or wm:
            if not hq and not wm:
                loc_str = loc
            elif not hq and not loc:
                loc_str = wm
            elif not loc and not wm:
                loc_str = hq
            elif not hq:
                loc_str = f"{loc} ({wm})"
            elif not loc:
                loc_str = f"{wm} ({hq})"
            elif not wm:
                loc_str = f"{loc} ({hq})"
            else:
                loc_str = f"{loc} ({wm}, {hq})"
            _run(tp, sep, font=body_font, size=typo["body_size"])
            _run(tp, loc_str, font=body_font, size=typo["body_size"])

        if role.get("date_start") or role.get("date_end"):
            tp.add_run("\t")
            _run(tp, f"{role.get('date_start', '')} – {role.get('date_end', '')}",
                 font=body_font, size=typo["body_size"])

        # Company description
        if exp_cfg.get("show_company_description", True) and role.get("company_description"):
            dp = doc.add_paragraph()
            _set_spacing(dp, after=40)
            _run(dp, role["company_description"].strip(), font=body_font,
                 size=typo["body_size"], italic=True)

        # Bullets
        mx = exp_cfg.get("max_bullets_per_role", 6)
        for b in role.get("bullets", [])[:mx]:
            bp = doc.add_paragraph(style="List Bullet")
            _set_spacing(bp, after=0, line=line_sp)
            _run(bp, b.strip(), font=body_font, size=typo["body_size"])


def _render_education(doc, data, cfg, usable_w):
    if not data.get("education"):
        return
    _heading(doc, "Education", cfg, key="education")
    typo = cfg["typography"]
    body_font = typo.get("body_font", typo.get("font", "Calibri"))
    sep = cfg.get("education", {}).get("separator", " | ")

    for i, edu in enumerate(data["education"]):
        dp = doc.add_paragraph()
        _set_spacing(dp, before=60 if i else 0, after=0)
        dp.paragraph_format.tab_stops.add_tab_stop(usable_w, WD_TAB_ALIGNMENT.RIGHT)

        status = f" ({edu['status']})" if edu.get("status") else ""
        _run(dp, f"{edu['degree']}{status}", font=body_font,
             size=typo["body_size"], bold=True)

        # Inline institution and location after degree
        extra = []
        if edu.get("institution"):
            extra.append(edu["institution"])
        if edu.get("location"):
            extra.append(edu["location"])
        if extra:
            _run(dp, sep, font=body_font, size=typo["body_size"])
            _run(dp, sep.join(extra), font=body_font, size=typo["body_size"])

        if edu.get("date_start") or edu.get("date_end"):
            dp.add_run("\t")
            _run(dp, f"{edu.get('date_start', '')} – {edu.get('date_end', '')}",
                 font=body_font, size=typo["body_size"])

        if edu.get("gpa"):
            gp = doc.add_paragraph()
            _set_spacing(gp, after=0)
            _run(gp, f"GPA: {edu['gpa']}", font=body_font, size=typo["body_size"])


def _render_certifications(doc, data, cfg, usable_w):
    if not data.get("certifications"):
        return
    _heading(doc, "Certifications", cfg, key="certifications")
    typo = cfg["typography"]
    body_font = typo.get("body_font", typo.get("font", "Calibri"))
    link_hex = cfg["colors"]["link"]
    sep = cfg.get("certifications", {}).get("separator", " | ")
    for i, cert in enumerate(data["certifications"]):
        p = doc.add_paragraph()
        _set_spacing(p, after=20, line=int(typo["line_spacing"] * 240))
        p.paragraph_format.tab_stops.add_tab_stop(usable_w, WD_TAB_ALIGNMENT.RIGHT)
        _run(p, cert["name"], font=body_font, size=typo["body_size"], bold=True)
        # Inline issuer and verify link
        _run(p, sep, font=body_font, size=typo["body_size"])
        _run(p, cert["issuer"], font=body_font, size=typo["body_size"])
        if cert.get("url"):
            _run(p, sep, font=body_font, size=typo["body_size"])
            _add_hyperlink(p, cert["url"], "Verify", link_hex,
                           font=body_font, size=typo["body_size"])
        if cert.get("date"):
            p.add_run("\t")
            _run(p, cert["date"], font=body_font, size=typo["body_size"])


def _render_languages(doc, data, cfg, _w):
    if not data.get("languages"):
        return
    _heading(doc, "Languages", cfg, key="languages")
    typo = cfg["typography"]
    body_font = typo.get("body_font", typo.get("font", "Calibri"))
    sep = cfg.get("languages", {}).get("separator", " · ")
    p = doc.add_paragraph()
    _set_spacing(p, after=0)
    for i, lang in enumerate(data["languages"]):
        if i > 0:
            _run(p, sep, font=body_font, size=typo["body_size"])
        _run(p, f"{lang['language']}:", font=body_font, size=typo["body_size"], bold=True)
        _run(p, f" {lang['level']}", font=body_font, size=typo["body_size"])


# ---------------------------------------------------------------------------
# Generic section renderers (for style-defined custom sections)
# ---------------------------------------------------------------------------

def _render_inline_section(doc, data, cfg, key, heading, _w):
    """Render a string-array section as items joined with a separator on one line."""
    items = data.get(key)
    if not items:
        return
    _heading(doc, heading, cfg, key=key)
    typo = cfg["typography"]
    body_font = typo.get("body_font", typo.get("font", "Calibri"))
    sep = cfg.get(key, {}).get("separator", " · ")
    p = doc.add_paragraph()
    _set_spacing(p, after=0)
    _run(p, sep.join(items), font=body_font, size=typo["body_size"])


def _render_list_section(doc, data, cfg, key, heading, _w):
    """Render a string-array section as bullet points."""
    items = data.get(key)
    if not items:
        return
    _heading(doc, heading, cfg, key=key)
    typo = cfg["typography"]
    body_font = typo.get("body_font", typo.get("font", "Calibri"))
    line_sp = int(typo["line_spacing"] * 240)
    for item in items:
        bp = doc.add_paragraph(style="List Bullet")
        _set_spacing(bp, after=0, line=line_sp)
        _add_inline_text(bp, item.strip(), font=body_font, size=typo["body_size"])


def _render_paragraph_section(doc, data, cfg, key, heading, _w):
    """Render a string-array section as plain paragraphs (no bullets)."""
    items = data.get(key)
    if not items:
        return
    _heading(doc, heading, cfg, key=key)
    typo = cfg["typography"]
    body_font = typo.get("body_font", typo.get("font", "Calibri"))
    line_sp = int(typo["line_spacing"] * 240)
    for item in items:
        pp = doc.add_paragraph()
        _set_spacing(pp, after=60, line=line_sp)
        _add_inline_text(pp, item.strip(), font=body_font, size=typo["body_size"])


def _render_structured_section(doc, data, cfg, key, heading, _w):
    """Render a section with structured (dict) items as formatted entries."""
    items = data.get(key)
    if not items:
        return
    _heading(doc, heading, cfg, key=key)
    typo = cfg["typography"]
    body_font = typo.get("body_font", typo.get("font", "Calibri"))
    line_sp = int(typo["line_spacing"] * 240)
    for item in items:
        pp = doc.add_paragraph()
        _set_spacing(pp, after=60, line=line_sp)
        if isinstance(item, dict):
            values = [str(v) for v in item.values() if v]
            if values:
                # Bold the first field (name/title), rest as regular text
                _run(pp, values[0], font=body_font, size=typo["body_size"], bold=True)
                if len(values) > 1:
                    _run(pp, " \u2014 " + " \u2014 ".join(values[1:]),
                         font=body_font, size=typo["body_size"])
        else:
            _add_inline_text(pp, str(item).strip(), font=body_font, size=typo["body_size"])


# Map config section names → renderer functions
_RENDERERS = {
    "summary": _render_summary,
    "skills": _render_skills,
    "experience": _render_experience,
    "education": _render_education,
    "certifications": _render_certifications,
    "languages": _render_languages,
}


# ---------------------------------------------------------------------------
# Main entry
# ---------------------------------------------------------------------------

def render_cv(data, config, output_path):
    """Build a CV .docx from structured *data* using *config* settings."""
    doc = Document()

    # Reset List Bullet style spacing to prevent extra gaps
    bullet_style = doc.styles['List Bullet']
    bullet_style.paragraph_format.space_before = Pt(0)
    bullet_style.paragraph_format.space_after = Pt(0)

    # ── Page setup ──
    section = doc.sections[0]
    pg = config["page"]
    if pg["size"] == "A4":
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    else:  # Letter
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.94)
    m = pg["margins"]
    section.top_margin = Cm(m["top"])
    section.bottom_margin = Cm(m["bottom"])
    section.left_margin = Cm(m["left"])
    section.right_margin = Cm(m["right"])

    usable_w = section.page_width - section.left_margin - section.right_margin
    typo = config["typography"]
    heading_font = typo.get("heading_font", typo.get("font", "Calibri"))
    body_font = typo.get("body_font", typo.get("font", "Calibri"))
    hdr_cfg = config["header"]
    hdr = data["header"]
    primary = hex_to_rgb(hdr_cfg.get("name_color", config["colors"]["primary"]))
    link_hex = config["colors"]["link"]
    hdr_align = (WD_ALIGN_PARAGRAPH.LEFT if hdr_cfg.get("alignment", "center") == "left"
                 else WD_ALIGN_PARAGRAPH.CENTER)

    # ── Header ──
    # Check if a photo should be included in the header
    photo_cfg = hdr_cfg.get("photo", {})
    photo_path_raw = hdr.get("photo_path")
    if photo_path_raw:
        photo_file = Path(photo_path_raw)
        if not photo_file.is_absolute():
            project_root_candidate = Path.cwd() / photo_file
            output_dir_candidate = Path(output_path).parent / photo_file
            if project_root_candidate.exists():
                photo_file = project_root_candidate
            else:
                photo_file = output_dir_candidate
    else:
        photo_file = None
    use_photo = (photo_cfg.get("enabled", False)
                 and photo_file is not None
                 and photo_file.exists())

    if photo_cfg.get("enabled", False) and not use_photo:
        missing = photo_path_raw or "(not set)"
        print(f"  Warning: Photo enabled in config but file not found: {missing}")

    if use_photo:
        # Table-based header: left column for text, right column for photo
        header_tbl = doc.add_table(rows=1, cols=2)
        header_tbl.autofit = False
        _remove_table_borders(header_tbl)
        pw = Cm(photo_cfg.get("width", 3.5))
        ph = Cm(photo_cfg.get("height", 4.5))
        pcol = pw + Cm(0.3)
        header_tbl.columns[0].width = usable_w - pcol
        header_tbl.columns[1].width = pcol
        # Photo in right cell
        right_cell = header_tbl.cell(0, 1)
        photo_p = right_cell.paragraphs[0]
        photo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(photo_p, before=0, after=0)
        photo_p.add_run().add_picture(str(photo_file), width=pw, height=ph)
        # Text content in left cell — use existing first paragraph for name
        left_cell = header_tbl.cell(0, 0)
        np = left_cell.paragraphs[0]
        _add_header_p = left_cell.add_paragraph
    else:
        np = doc.add_paragraph()
        _add_header_p = doc.add_paragraph

    # Name
    np.alignment = hdr_align
    _set_spacing(np, after=0)
    name_text = (hdr["name"].upper() if hdr_cfg["name_case"] == "upper"
                 else hdr["name"].title() if hdr_cfg["name_case"] == "title"
                 else hdr["name"])
    _run(np, name_text, font=heading_font, size=hdr_cfg["name_font_size"],
         bold=hdr_cfg.get("name_bold", True), color=primary)

    # Titles
    if hdr_cfg.get("show_titles", True) and hdr.get("titles"):
        tp = _add_header_p()
        tp.alignment = hdr_align
        _set_spacing(tp, after=0)
        _run(tp, " · ".join(hdr["titles"]), font=body_font,
             size=hdr_cfg.get("titles_font_size", typo["body_size"]))

    # Contact line
    contact = []
    for key in ("location", "phone", "email"):
        if hdr.get(key):
            contact.append(hdr[key])
    if contact:
        cp = _add_header_p()
        cp.alignment = hdr_align
        _set_spacing(cp, after=0)
        _run(cp, " | ".join(contact), font=body_font,
             size=hdr_cfg.get("contact_font_size", typo["body_size"]))

    # Links line
    has_links = hdr.get("linkedin") or hdr.get("github")
    if has_links:
        lp = _add_header_p()
        lp.alignment = hdr_align
        _set_spacing(lp, after=60)
        links_size = hdr_cfg.get("links_font_size", typo["body_size"])
        if hdr.get("linkedin"):
            display = hdr["linkedin"].replace("https://", "").replace("http://", "").rstrip("/")
            _add_hyperlink(lp, hdr["linkedin"], display, link_hex, font=body_font, size=links_size)
        if hdr.get("linkedin") and hdr.get("github"):
            _run(lp, " | ", font=body_font, size=links_size)
        if hdr.get("github"):
            display = hdr["github"].replace("https://", "").replace("http://", "").rstrip("/")
            _add_hyperlink(lp, hdr["github"], display, link_hex, font=body_font, size=links_size)

    # Relocation
    if hdr_cfg.get("show_relocation", True) and hdr.get("relocation"):
        rp = _add_header_p()
        rp.alignment = hdr_align
        _set_spacing(rp, after=60)
        _run(rp, hdr["relocation"], font=body_font,
             size=hdr_cfg.get("relocation_font_size", typo["body_size"]),
             bold=hdr_cfg.get("relocation_bold", True))

    # ── Sections ──
    section_types = config.get("sections", {}).get("section_types", {})
    for name in config["sections"]["order"]:
        renderer = _RENDERERS.get(name)
        if renderer:
            renderer(doc, data, config, usable_w)
        elif name in data:
            # Generic rendering for style-defined custom sections
            value = data[name]
            heading = name.replace("_", " ").title()
            if isinstance(value, str):
                # Single string section (e.g., self_pr, objective)
                _heading(doc, heading, config, key=name)
                p = doc.add_paragraph()
                _set_spacing(p, after=60, line=int(typo["line_spacing"] * 240))
                _run(p, value.strip(), font=body_font, size=typo["body_size"])
            elif isinstance(value, list) and value and isinstance(value[0], dict):
                # Structured object items (e.g., references, volunteer_work)
                _render_structured_section(doc, data, config, name, heading, usable_w)
            else:
                # String array — use configured layout
                layout = section_types.get(name, "list")
                if layout == "inline":
                    _render_inline_section(doc, data, config, name, heading, usable_w)
                elif layout == "paragraph":
                    _render_paragraph_section(doc, data, config, name, heading, usable_w)
                else:
                    _render_list_section(doc, data, config, name, heading, usable_w)

    doc.save(str(output_path))
    return str(output_path)
